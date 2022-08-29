from multiprocessing.resource_sharer import stop
from docx import Document
from docx.shared import Inches

document = Document()

#profile picture
document.add_picture('Abbah.jpg', width =Inches(2.0))

#name, phone_number and email
name = input ('what is your name?')
phone_number = input ('what is your phone number?')
email = input ('what is your email?')

document.add_paragraph(
    name +'| ' + phone_number + '| ' + email
)

#about me 

document.add_heading('About me')
about_me = input('tell me about yourself')

#work experience

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input ('Enter Company')
from_date = input('From Date')
to_date = input ('To Date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' ' + to_date + '\n').italic = True

experience_details = input (
    'Describe your experience at ' + company)
p.add_run(experience_details)


#more experiences
while True:
    has_more_experiences = input(
        'Do yo have more experiences? Yes or No')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

    company = input ('Enter Company')
    from_date = input('From Date')
    to_date = input ('To Date')

    p.add_run(company + ' ').bold = True
    p.add_run(from_date + ' ' + to_date + '\n').italic = True

    experience_details = input (
        'Describe your experience at ' + company)
    p.add_run(experience_details)

#TODO
document.save('cv.docx')
